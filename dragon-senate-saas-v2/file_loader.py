"""
LobeHub-inspired business file loader for PDF/Word/Excel/text.
"""

from __future__ import annotations

import io
import re
from dataclasses import asdict, dataclass, field
from pathlib import Path
from typing import Any


@dataclass(slots=True)
class LoadedFile:
    filename: str
    file_type: str
    raw_text: str
    metadata: dict[str, Any] = field(default_factory=dict)
    structured_data: dict[str, Any] = field(default_factory=dict)
    extraction_quality: float = 1.0

    def to_dict(self) -> dict[str, Any]:
        return asdict(self)


@dataclass(slots=True)
class BusinessCardExtract:
    name: str | None = None
    title: str | None = None
    company: str | None = None
    phone: str | None = None
    email: str | None = None
    wechat: str | None = None
    address: str | None = None
    raw_text: str | None = None

    def to_dict(self) -> dict[str, Any]:
        return asdict(self)


class LobsterFileLoader:
    SUPPORTED_TYPES = {".pdf", ".docx", ".doc", ".xlsx", ".xls", ".txt", ".md", ".csv"}

    async def load(self, file_path: str, file_bytes: bytes | None = None) -> LoadedFile:
        path = Path(file_path)
        ext = path.suffix.lower()
        if ext not in self.SUPPORTED_TYPES:
            raise ValueError(f"unsupported file type: {ext}")
        if ext == ".pdf":
            return await self._load_pdf(path, file_bytes)
        if ext in {".docx", ".doc"}:
            return await self._load_docx(path, file_bytes)
        if ext in {".xlsx", ".xls", ".csv"}:
            return await self._load_excel(path, file_bytes, ext)
        return await self._load_text(path, file_bytes)

    async def extract_business_card(self, file: LoadedFile) -> BusinessCardExtract:
        text = str(file.raw_text or "").strip()
        if not text:
            return BusinessCardExtract(raw_text="")
        return self._rule_based_extract(text)

    async def extract_leads_from_excel(self, file: LoadedFile) -> list[dict[str, Any]]:
        headers = list(file.structured_data.get("headers") or [])
        rows = list(file.structured_data.get("rows") or [])
        results: list[dict[str, Any]] = []
        for row in rows:
            if isinstance(row, dict):
                results.append(dict(row))
                continue
            if isinstance(row, list) and headers:
                results.append({headers[idx]: row[idx] if idx < len(row) else "" for idx in range(len(headers))})
        return results

    async def _load_pdf(self, path: Path, file_bytes: bytes | None) -> LoadedFile:
        text = ""
        metadata: dict[str, Any] = {}
        try:
            try:
                from pypdf import PdfReader  # type: ignore
            except Exception:
                from PyPDF2 import PdfReader  # type: ignore

            reader = PdfReader(io.BytesIO(file_bytes)) if file_bytes is not None else PdfReader(str(path))
            pages = []
            for page in reader.pages:
                try:
                    pages.append(page.extract_text() or "")
                except Exception:
                    pages.append("")
            text = "\n".join(pages).strip()
            metadata = {
                "page_count": len(reader.pages),
                "title": str(getattr(reader.metadata, "title", "") or ""),
                "author": str(getattr(reader.metadata, "author", "") or ""),
            }
            quality = 1.0 if text else 0.4
            return LoadedFile(filename=path.name, file_type="pdf", raw_text=text, metadata=metadata, extraction_quality=quality)
        except Exception as exc:
            return LoadedFile(
                filename=path.name,
                file_type="pdf",
                raw_text="",
                metadata={"error": str(exc)},
                extraction_quality=0.0,
            )

    async def _load_docx(self, path: Path, file_bytes: bytes | None) -> LoadedFile:
        try:
            import docx  # type: ignore

            doc = docx.Document(io.BytesIO(file_bytes)) if file_bytes is not None else docx.Document(str(path))
            paragraphs = [para.text.strip() for para in doc.paragraphs if para.text.strip()]
            text = "\n".join(paragraphs)
            metadata = {"paragraph_count": len(paragraphs)}
            return LoadedFile(filename=path.name, file_type="docx", raw_text=text, metadata=metadata, extraction_quality=1.0 if text else 0.5)
        except Exception as exc:
            return LoadedFile(filename=path.name, file_type="docx", raw_text="", metadata={"error": str(exc)}, extraction_quality=0.0)

    async def _load_excel(self, path: Path, file_bytes: bytes | None, ext: str) -> LoadedFile:
        try:
            if ext == ".csv":
                raw = (file_bytes or path.read_bytes()).decode("utf-8", errors="replace")
                rows = [line.split(",") for line in raw.splitlines() if line.strip()]
                headers = rows[0] if rows else []
                body = rows[1:] if len(rows) > 1 else []
                text = "\n".join(", ".join(cell for cell in row) for row in rows[:50])
                return LoadedFile(
                    filename=path.name,
                    file_type="csv",
                    raw_text=text,
                    metadata={"row_count": len(body), "column_count": len(headers)},
                    structured_data={"headers": headers, "rows": body},
                    extraction_quality=1.0,
                )

            from openpyxl import load_workbook  # type: ignore

            workbook = load_workbook(io.BytesIO(file_bytes), data_only=True) if file_bytes is not None else load_workbook(str(path), data_only=True)
            sheet = workbook.active
            all_rows = list(sheet.iter_rows(values_only=True))
            headers = [str(cell or "").strip() for cell in (all_rows[0] if all_rows else [])]
            rows = [[cell for cell in row] for row in all_rows[1:]]
            text_rows = [" | ".join(str(cell or "") for cell in row) for row in all_rows[:50]]
            return LoadedFile(
                filename=path.name,
                file_type="xlsx",
                raw_text="\n".join(text_rows),
                metadata={"sheet": sheet.title, "row_count": len(rows), "column_count": len(headers)},
                structured_data={"headers": headers, "rows": rows},
                extraction_quality=1.0 if rows else 0.6,
            )
        except Exception as exc:
            return LoadedFile(filename=path.name, file_type="excel", raw_text="", metadata={"error": str(exc)}, extraction_quality=0.0)

    async def _load_text(self, path: Path, file_bytes: bytes | None) -> LoadedFile:
        raw = file_bytes if file_bytes is not None else path.read_bytes()
        text = self._decode_text(raw)
        return LoadedFile(
            filename=path.name,
            file_type=path.suffix.lower().lstrip(".") or "txt",
            raw_text=text,
            metadata={"length": len(text)},
            extraction_quality=1.0 if text else 0.0,
        )

    @staticmethod
    def _decode_text(payload: bytes) -> str:
        for encoding in ("utf-8", "utf-8-sig", "gb18030", "latin-1"):
            try:
                text = payload.decode(encoding)
                if text:
                    return text
            except Exception:
                continue
        return payload.decode("utf-8", errors="replace")

    def _rule_based_extract(self, text: str) -> BusinessCardExtract:
        lines = [line.strip() for line in re.split(r"[\n\r]+", text) if line.strip()]
        phone = _first_match(r"\b1[3-9]\d{9}\b", text)
        email = _first_match(r"[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}", text)
        wechat = _first_match(r"(?:微信|wechat|vx)[:： ]*([A-Za-z0-9_-]{4,32})", text, group=1)
        title = None
        company = None
        name = None
        title_keywords = ("CEO", "CTO", "CFO", "COO", "VP", "Director", "Manager", "总监", "经理", "顾问", "创始人", "负责人", "老师", "医生")
        company_keywords = ("公司", "科技", "集团", "工作室", "诊所", "美容院", "机构", "Tech", "Inc", "Corp", "LLC", "Ltd")
        for line in lines[:8]:
            if name is None and 1 <= len(line) <= 32 and not any(ch.isdigit() for ch in line) and "@" not in line and len(line.split()) <= 4:
                name = line
            if title is None and any(token.lower() in line.lower() for token in title_keywords):
                title = line
            if company is None and any(token in line for token in company_keywords):
                company = line
        return BusinessCardExtract(
            name=name,
            title=title,
            company=company,
            phone=phone,
            email=email,
            wechat=wechat,
            address=None,
            raw_text=text,
        )


def _first_match(pattern: str, text: str, *, group: int = 0) -> str | None:
    match = re.search(pattern, text, re.IGNORECASE)
    if not match:
        return None
    return str(match.group(group)).strip()
